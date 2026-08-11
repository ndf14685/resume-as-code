"""ATS-safe DOCX renderer (python-docx).

Deliberately simple: single column, one standard font (Calibri), real text and
real bullets, no images, no icons, no text boxes, and NO tables used for
layout. Company / title / dates are plain paragraph text so any ATS can parse
them. Right-aligned dates use a tab stop (text), not a table cell.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from .layout import contact_line, language_line, present_sections, training_line
from .models import ResumeModel

FONT = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x44, 0x44, 0x44)
ACCENT = RGBColor(0x1F, 0x2D, 0x3D)


def _set_margins(section, inches: float) -> None:
    section.top_margin = Inches(inches)
    section.bottom_margin = Inches(inches)
    section.left_margin = Inches(inches)
    section.right_margin = Inches(inches)


def _usable_width(section):
    return section.page_width - section.left_margin - section.right_margin


def _tight(p, before=0, after=2):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.06


def _run(p, text, *, size=10.5, bold=False, italic=False, color=INK, caps=False):
    r = p.add_run(text.upper() if caps else text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def _bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "9AA5B1")
    borders.append(bottom)
    pPr.append(borders)


def _heading(doc, text):
    p = doc.add_paragraph()
    _tight(p, before=9, after=3)
    _run(p, text, size=11.5, bold=True, color=ACCENT, caps=True)
    _bottom_border(p)
    return p


def _title_with_dates(doc, section, left_text, right_text, *, bold=True, size=10.5):
    p = doc.add_paragraph()
    _tight(p, before=5, after=0)
    p.paragraph_format.tab_stops.add_tab_stop(
        _usable_width(section), WD_TAB_ALIGNMENT.RIGHT
    )
    _run(p, left_text, size=size, bold=bold)
    _run(p, "\t" + right_text, size=size - 0.5, color=MUTED)
    return p


def render_docx(resume: ResumeModel, out_path: str | Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    _set_margins(section, 0.6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    # ---- Header -----------------------------------------------------------
    name_p = doc.add_paragraph()
    _tight(name_p, before=0, after=1)
    _run(name_p, resume.name, size=21, bold=True, color=ACCENT)

    head_p = doc.add_paragraph()
    _tight(head_p, before=0, after=1)
    _run(head_p, resume.headline, size=11.5, bold=True, color=INK)
    if resume.tagline:
        _run(head_p, "   " + resume.tagline, size=10, color=MUTED)

    contact_p = doc.add_paragraph()
    _tight(contact_p, before=0, after=2)
    _run(contact_p, contact_line(resume), size=9.5, color=MUTED)

    # ---- Sections ---------------------------------------------------------
    for key, title in present_sections(resume):
        _heading(doc, title)
        if key == "summary":
            p = doc.add_paragraph()
            _tight(p, before=2, after=2)
            _run(p, resume.summary, size=10.5)
        elif key == "skills":
            for g in resume.skill_groups:
                p = doc.add_paragraph()
                _tight(p, before=0, after=1)
                _run(p, f"{g.name}: ", bold=True, size=10)
                _run(p, ", ".join(g.items), size=10)
        elif key == "experience":
            for exp in resume.experiences:
                if exp.condensed:
                    _title_with_dates(
                        doc, section,
                        f"{exp.title}, {exp.company}",
                        exp.meta_right, bold=True, size=10,
                    )
                    continue
                _title_with_dates(doc, section, exp.title, exp.meta_right)
                sub = " · ".join(
                    x for x in [exp.company, exp.engagement_label, exp.location] if x
                )
                sp = doc.add_paragraph()
                _tight(sp, before=0, after=1)
                _run(sp, sub, size=10, italic=True, color=MUTED)
                for b in exp.bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    _tight(bp, before=0, after=1)
                    bp.paragraph_format.left_indent = Inches(0.22)
                    _run(bp, b, size=10)
        elif key == "projects":
            for proj in resume.featured_projects:
                _title_with_dates(doc, section, proj.name, proj.label)
                if proj.tagline:
                    tp = doc.add_paragraph()
                    _tight(tp, before=0, after=1)
                    _run(tp, proj.tagline, size=10, italic=True, color=MUTED)
                for b in proj.bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    _tight(bp, before=0, after=1)
                    bp.paragraph_format.left_indent = Inches(0.22)
                    _run(bp, b, size=10)
        elif key == "certifications":
            for c in resume.certifications:
                p = doc.add_paragraph()
                _tight(p)
                _run(p, " · ".join(str(v) for v in c.values() if v), size=10)
        elif key == "training":
            for g in resume.training:
                p = doc.add_paragraph()
                _tight(p, before=0, after=1)
                _run(p, f"{g.get('provider', '')}: ", bold=True, size=10)
                items = []
                for it in g.get("items", []):
                    yr = it.get("year")
                    items.append(f"{it.get('name','')} ({yr})" if yr else it.get("name", ""))
                _run(p, ", ".join(items), size=10)
        elif key == "languages":
            for item in resume.languages:
                p = doc.add_paragraph()
                _tight(p, before=0, after=1)
                _run(p, f"{item.get('language','')}: ", bold=True, size=10)
                _run(p, item.get("level", ""), size=10)
        elif key == "education":
            for e in resume.education:
                p = doc.add_paragraph()
                _tight(p)
                _run(p, " · ".join(str(v) for v in e.values() if v), size=10)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
